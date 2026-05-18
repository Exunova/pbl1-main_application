#!/usr/bin/env python3
import os
import re
import sys
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
EVIDENCE_DIR = os.path.join(PROJECT_ROOT, ".sisyphus", "evidence")
SCREENSHOT_DIR = os.path.join(EVIDENCE_DIR, "task-23")
OUTPUT_FILE = os.path.join(PROJECT_ROOT, "MAPRO_Test_Report.docx")
TEST_CASES_MD = os.path.join(PROJECT_ROOT, "MAPRO_Test_Cases.md")

TEST_RESULT_FILES = {
    "SCR": os.path.join(EVIDENCE_DIR, "task-20-scraper-tests.txt"),
    "CAC": os.path.join(EVIDENCE_DIR, "task-21-caching-tests.txt"),
    "IPC": os.path.join(EVIDENCE_DIR, "task-22-ipc-tests.txt"),
}

SCREENSHOT_FILES = {
    "SCR": os.path.join(SCREENSHOT_DIR, "task-23-scraper-tests.png"),
    "CAC": os.path.join(SCREENSHOT_DIR, "task-23-caching-tests.png"),
    "IPC": os.path.join(SCREENSHOT_DIR, "task-23-ipc-tests.png"),
}


def parse_test_results(filepath):
    unique_results = {}
    if not os.path.exists(filepath):
        print(f"Warning: {filepath} not found")
        return []

    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    pattern = r"test_([a-z]{3})_(\d{3})_[^\s]+\s+(PASSED|FAILED|SKIPPED|ERROR)"
    for match in re.finditer(pattern, content, re.IGNORECASE):
        prefix = match.group(1).upper()
        num = match.group(2)
        result = match.group(3).upper()
        test_id = f"{prefix}-{num}"
        if test_id not in unique_results:
            unique_results[test_id] = {"id": test_id, "result": result}

    return list(unique_results.values())


def parse_test_descriptions(md_path):
    descriptions = {}
    if not os.path.exists(md_path):
        print(f"Warning: {md_path} not found")
        return descriptions

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    pattern = r"\|\s*([A-Z]{3}-\d{3})\s*\|\s*([^|]+)\|"
    for match in re.finditer(pattern, content):
        test_id = match.group(1).strip()
        desc = match.group(2).strip()
        descriptions[test_id] = desc

    return descriptions


def add_styled_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x20, 0x37)
    return heading


def add_styled_paragraph(doc, text, bold=False, italic=False, font_size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_cover_page(doc):
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MAPRO Test Report")
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x20, 0x37)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Multimarket Analytics Portfolio Tracker")
    run.font.size = Pt(18)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version_para.add_run("Version: 1.0")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_page_break()


def add_executive_summary(doc, all_results):
    add_styled_heading(doc, "Executive Summary", level=1)

    total = len(all_results)
    passed = sum(1 for r in all_results if r["result"] == "PASSED")
    failed = sum(1 for r in all_results if r["result"] == "FAILED")
    skipped = sum(1 for r in all_results if r["result"] == "SKIPPED")
    errors = sum(1 for r in all_results if r["result"] == "ERROR")

    add_styled_paragraph(doc, f"This report summarizes the test execution results for the MAPRO application.")
    add_styled_paragraph(doc, f"A total of {total} test cases were executed across three modules: Scraping Engine, Caching & TTL, and IPC.")

    doc.add_paragraph()

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    summary_data = [
        ("Total Tests", str(total)),
        ("Passed", str(passed)),
        ("Failed", str(failed)),
        ("Skipped", str(skipped)),
        ("Errors", str(errors)),
    ]

    for i, (label, value) in enumerate(summary_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.bold = True if i == 0 else False

    doc.add_paragraph()

    add_styled_heading(doc, "Environment Information", level=2)
    add_styled_paragraph(doc, f"Platform: Darwin (macOS)")
    add_styled_paragraph(doc, f"Python: 3.14.4")
    add_styled_paragraph(doc, f"pytest: 9.0.3")
    add_styled_paragraph(doc, f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_page_break()

    return {"total": total, "passed": passed, "failed": failed, "skipped": skipped, "errors": errors}


def add_test_section(doc, section_title, test_results, descriptions, screenshot_path):
    add_styled_heading(doc, section_title, level=1)

    if screenshot_path and os.path.exists(screenshot_path):
        add_styled_paragraph(doc, "Screenshot:", bold=True)
        try:
            doc.add_picture(screenshot_path, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            add_styled_paragraph(doc, f"[Error embedding screenshot: {e}]")
        doc.add_paragraph()
    else:
        add_styled_paragraph(doc, "[Screenshot not available]")
        doc.add_paragraph()

    add_styled_paragraph(doc, f"Total test cases in this section: {len(test_results)}", bold=True)
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    headers = ["Test ID", "Description", "Result", "Screenshot Reference"]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    for result in test_results:
        test_id = result["id"]
        result_status = result["result"]
        description = descriptions.get(test_id, "No description available")
        screenshot_ref = os.path.basename(screenshot_path) if screenshot_path and os.path.exists(screenshot_path) else "N/A"

        row_cells = table.add_row().cells
        row_cells[0].text = test_id
        row_cells[1].text = description
        row_cells[2].text = result_status
        row_cells[3].text = screenshot_ref

        for paragraph in row_cells[2].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                if result_status == "PASSED":
                    run.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)
                elif result_status == "FAILED":
                    run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
                elif result_status == "SKIPPED":
                    run.font.color.rgb = RGBColor(0xF5, 0xA6, 0x23)
                elif result_status == "ERROR":
                    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    doc.add_page_break()


def generate_report():
    print("=" * 60)
    print("MAPRO Test Report Generator")
    print("=" * 60)

    print("Parsing test descriptions from MAPRO_Test_Cases.md...")
    descriptions = parse_test_descriptions(TEST_CASES_MD)
    print(f"  Found {len(descriptions)} test descriptions")

    all_results = []
    section_results = {}

    for prefix, filepath in TEST_RESULT_FILES.items():
        print(f"Parsing test results from {os.path.basename(filepath)}...")
        results = parse_test_results(filepath)
        section_results[prefix] = results
        all_results.extend(results)
        print(f"  Found {len(results)} test results")

    print("Creating Word document...")
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    add_cover_page(doc)
    summary = add_executive_summary(doc, all_results)

    add_test_section(
        doc,
        "Section 1: Scraping Engine",
        section_results.get("SCR", []),
        descriptions,
        SCREENSHOT_FILES.get("SCR")
    )

    add_test_section(
        doc,
        "Section 2: Caching & TTL",
        section_results.get("CAC", []),
        descriptions,
        SCREENSHOT_FILES.get("CAC")
    )

    add_test_section(
        doc,
        "Section 3: IPC (Inter-Process Communication)",
        section_results.get("IPC", []),
        descriptions,
        SCREENSHOT_FILES.get("IPC")
    )

    add_styled_heading(doc, "Additional Screenshots", level=1)
    add_styled_paragraph(doc, "The following screenshots were captured during test execution:")
    doc.add_paragraph()

    extra_screenshots = [
        ("Loading Screen", os.path.join(SCREENSHOT_DIR, "task-23-loading-screen.png")),
        ("New Data Badge", os.path.join(SCREENSHOT_DIR, "task-23-new-data-badge.png")),
    ]

    for label, path in extra_screenshots:
        if os.path.exists(path):
            add_styled_paragraph(doc, f"{label}:", bold=True)
            try:
                doc.add_picture(path, width=Inches(5.0))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                add_styled_paragraph(doc, f"[Error embedding screenshot: {e}]")
            doc.add_paragraph()

    add_styled_heading(doc, "End of Report", level=1)
    add_styled_paragraph(doc, f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    add_styled_paragraph(doc, "This document was automatically generated by tests/generate_report.py")

    doc.save(OUTPUT_FILE)
    print(f"\nReport saved to: {OUTPUT_FILE}")

    file_size = os.path.getsize(OUTPUT_FILE)
    print(f"File size: {file_size} bytes ({file_size / 1024:.1f} KB)")

    if file_size < 10 * 1024:
        print("WARNING: File size is less than 10KB!")
        return False

    print("\nVerifying document content...")
    verify_doc = Document(OUTPUT_FILE)
    paragraph_count = len(verify_doc.paragraphs)
    print(f"  Paragraphs: {paragraph_count}")

    if paragraph_count <= 20:
        print("WARNING: Paragraph count is 20 or less!")
        return False

    scr_count = sum(1 for r in all_results if r["id"].startswith("SCR-"))
    cac_count = sum(1 for r in all_results if r["id"].startswith("CAC-"))
    ipc_count = sum(1 for r in all_results if r["id"].startswith("IPC-"))

    print(f"  SCR test IDs: {scr_count} (expected >= 40)")
    print(f"  CAC test IDs: {cac_count} (expected >= 19)")
    print(f"  IPC test IDs: {ipc_count} (expected >= 16)")

    if scr_count < 40 or cac_count < 19 or ipc_count < 16:
        print("WARNING: Not all test IDs found!")
        return False

    print("\nAll verifications passed!")
    return True


if __name__ == "__main__":
    success = generate_report()
    sys.exit(0 if success else 1)
